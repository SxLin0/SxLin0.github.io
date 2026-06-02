from html import escape
from pathlib import Path

from docx import Document


ARTICLE_SOURCES = {
    "software-major": "关于填报软件工程专业的一些看法.docx",
    "capital-scientific-thinking": "《资本论》中的科学思维 .docx",
    "digital-labor-alienation": "数字时代劳动异化的四重维度与解放路径.docx",
}

POEM_ESSAY_SOURCES = {
    "spring-essay": "浅谈“春”.docx",
}

SOURCE_DIR = Path("content/articles")
OUTPUT_DIR = Path("content/articles-html")
POEM_ESSAY_SOURCE_DIR = Path("content/poem-essay")
PRINT_OUTPUT_DIR = Path("content/print")


def paragraph_class(text, index):
    if index == 0:
        return "doc-title"
    if index == 1 and len(text) <= 40:
        return "doc-subtitle"
    if text.startswith(("摘要", "关键词")):
        return "doc-lead"
    if text.startswith("- "):
        return "doc-list-line"
    if text[:2] in {"一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、"}:
        return "doc-heading"
    if text[:2].isdigit() and text[2:3] in {"、", "."}:
        return "doc-heading"
    return "doc-paragraph"


def render_docx(source, output):
    doc = Document(source)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        raise ValueError(f"{source} has no readable paragraphs")

    title = paragraphs[0]
    body = [f"    <h1>{escape(title)}</h1>"]
    for index, text in enumerate(paragraphs[1:], start=1):
        cls = paragraph_class(text, index)
        tag = "h2" if cls == "doc-heading" else "p"
        body.append(f'    <{tag} class="{cls}">{escape(text)}</{tag}>')

    html = "\n".join([
        "<!DOCTYPE html>",
        '<html lang="zh-CN">',
        "<head>",
        '    <meta charset="UTF-8">',
        f"    <title>{escape(title)}</title>",
        "</head>",
        "<body>",
        *body,
        "</body>",
        "</html>",
        "",
    ])
    output.write_text(html, encoding="utf-8")


def render_essay_docx(source, output):
    doc = Document(source)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        raise ValueError(f"{source} has no readable paragraphs")

    title = paragraphs[0]
    signature = paragraphs[-2:] if len(paragraphs) > 2 else []
    body_paragraphs = paragraphs[1:-2] if signature else paragraphs[1:]
    body = [f"    <h1>{escape(title)}</h1>"]
    body.extend(f'    <p class="essay-paragraph">{escape(text)}</p>' for text in body_paragraphs)
    if signature:
        body.append('    <div class="signature">')
        body.extend(f"        <p>{escape(text)}</p>" for text in signature)
        body.append("    </div>")

    html = "\n".join([
        "<!DOCTYPE html>",
        '<html lang="zh-CN">',
        "<head>",
        '    <meta charset="UTF-8">',
        f"    <title>{escape(title)}</title>",
        "</head>",
        "<body>",
        *body,
        "</body>",
        "</html>",
        "",
    ])
    output.write_text(html, encoding="utf-8")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for slug, filename in ARTICLE_SOURCES.items():
        render_docx(SOURCE_DIR / filename, OUTPUT_DIR / f"{slug}.html")
    PRINT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for slug, filename in POEM_ESSAY_SOURCES.items():
        render_essay_docx(POEM_ESSAY_SOURCE_DIR / filename, PRINT_OUTPUT_DIR / f"{slug}.html")


if __name__ == "__main__":
    main()
